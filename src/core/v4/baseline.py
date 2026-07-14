"""Read-only paragraph baseline import and block intersection mapping."""

from __future__ import annotations

import hashlib
import json
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, List, Sequence, Tuple

from docx import Document

from .database import V4Database, stable_id, utc_now


@dataclass(frozen=True)
class SourceParagraph:
    index: int
    text: str
    char_start: int
    char_end: int


def paragraph_spans(text: str) -> List[SourceParagraph]:
    """Split exactly like the project paragraph model while retaining source offsets."""
    result: List[SourceParagraph] = []
    cursor = 0
    for delimiter in re.finditer(r"\n\s*\n", text):
        segment = text[cursor : delimiter.start()]
        stripped = segment.strip()
        if stripped:
            left = len(segment) - len(segment.lstrip())
            right = len(segment.rstrip())
            result.append(
                SourceParagraph(len(result), stripped, cursor + left, cursor + right)
            )
        cursor = delimiter.end()
    segment = text[cursor:]
    stripped = segment.strip()
    if stripped:
        left = len(segment) - len(segment.lstrip())
        right = len(segment.rstrip())
        result.append(SourceParagraph(len(result), stripped, cursor + left, cursor + right))
    return result


class DocxBaselineImporter:
    def __init__(self, database: V4Database, project):
        self.database = database
        self.project = project

    @staticmethod
    def _sha256(path: Path) -> str:
        digest = hashlib.sha256()
        with path.open("rb") as handle:
            for chunk in iter(lambda: handle.read(1024 * 1024), b""):
                digest.update(chunk)
        return digest.hexdigest()

    @staticmethod
    def _docx_paragraphs(path: Path) -> List[str]:
        document = Document(path)
        if document.tables:
            raise ValueError("第二版基线导入暂不接受含表格的DOCX")
        return [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]

    @staticmethod
    def _all_positions(source: str, needle: str) -> List[int]:
        positions: List[int] = []
        cursor = 0
        while True:
            found = source.find(needle, cursor)
            if found < 0:
                return positions
            positions.append(found)
            cursor = found + 1

    def import_docx(
        self,
        docx_path: str | Path,
        name: str = "legacy_docx",
        require_exact_count: bool = True,
    ) -> Dict[str, object]:
        path = Path(docx_path).resolve()
        if not path.exists():
            raise FileNotFoundError(f"DOCX不存在: {path}")
        source_text = self.project.source_file.read_text(encoding="utf-8")
        source_paragraphs = paragraph_spans(source_text)
        target_paragraphs = self._docx_paragraphs(path)
        if require_exact_count and len(source_paragraphs) != len(target_paragraphs):
            raise ValueError(
                "DOCX与英文源文的非空段落数不一致："
                f"source={len(source_paragraphs)}, docx={len(target_paragraphs)}"
            )
        aligned_count = min(len(source_paragraphs), len(target_paragraphs))
        file_hash = self._sha256(path)
        edition_id = self.database.active_source_edition_id()
        importer_version = "paragraph-v2"
        document_id = stable_id(
            "baseline", f"{importer_version}:{edition_id}:{name}:{file_hash}", length=24
        )
        blocks = self.database.list_blocks()
        block_positions: List[Tuple[object, int, str]] = []
        unaligned: List[str] = []
        ambiguous = 0
        position_cache: Dict[str, List[int]] = {}
        previous_start = -1
        previous_end = -1
        for block in blocks:
            positions = position_cache.setdefault(
                block.source_text,
                self._all_positions(source_text, block.source_text),
            )
            if not positions:
                unaligned.append(block.id)
                continue
            status = "exact"
            start = positions[0]
            if len(positions) > 1:
                ambiguous += 1
                status = "sequence_resolved_ambiguity"
                after_previous = [position for position in positions if position >= previous_end]
                overlapping = [position for position in positions if position > previous_start]
                if after_previous:
                    start = after_previous[0]
                elif overlapping:
                    start = overlapping[0]
                else:
                    status = "ambiguous_source_occurrence"
            elif positions[0] < previous_start:
                unaligned.append(block.id)
                continue
            block_positions.append((block, start, status))
            previous_start = start
            previous_end = start + len(block.source_text)
        if unaligned:
            preview = ", ".join(unaligned[:10])
            raise ValueError(f"有{len(unaligned)}个文本块无法在源文件中定位：{preview}")

        metadata = {
            "source_paragraph_count": len(source_paragraphs),
            "exact_count_required": require_exact_count,
            "importer_version": importer_version,
        }

        with self.database.transaction() as connection:
            existing = connection.execute(
                """SELECT * FROM baseline_documents
                   WHERE source_edition_id=? AND name=? AND file_sha256=?""",
                (edition_id, name, file_hash),
            ).fetchone()
            if existing:
                existing_metadata = json.loads(existing["metadata_json"] or "{}")
                linked_blocks = connection.execute(
                    "SELECT COUNT(DISTINCT block_id) FROM block_baseline_links WHERE baseline_document_id=?",
                    (existing["id"],),
                ).fetchone()[0]
                if (
                    existing_metadata.get("importer_version") == importer_version
                    and linked_blocks == len(blocks)
                ):
                    return {
                        "baseline_document_id": existing["id"],
                        "name": name,
                        "source_paragraphs": len(source_paragraphs),
                        "target_paragraphs": len(target_paragraphs),
                        "aligned_paragraphs": aligned_count,
                        "blocks": len(blocks),
                        "ambiguous_blocks": ambiguous,
                        "idempotent": True,
                        "file_sha256": file_hash,
                    }
                document_id = existing["id"]
                connection.execute(
                    "DELETE FROM block_baseline_links WHERE baseline_document_id=?",
                    (document_id,),
                )
                connection.execute(
                    "DELETE FROM baseline_paragraphs WHERE baseline_document_id=?",
                    (document_id,),
                )
                connection.execute(
                    """UPDATE baseline_documents SET file_path=?, paragraph_count=?,
                           metadata_json=?, active=1, created_at=? WHERE id=?""",
                    (
                        str(path),
                        len(target_paragraphs),
                        json.dumps(metadata, ensure_ascii=False),
                        utc_now(),
                        document_id,
                    ),
                )
            else:
                connection.execute(
                    "UPDATE baseline_documents SET active=0 WHERE source_edition_id=? AND name=?",
                    (edition_id, name),
                )
                connection.execute(
                    """INSERT INTO baseline_documents(
                           id, source_edition_id, name, kind, file_path, file_sha256,
                           paragraph_count, metadata_json, active, created_at
                       ) VALUES(?, ?, ?, 'docx_paragraph', ?, ?, ?, ?, 1, ?)""",
                    (
                        document_id,
                        edition_id,
                        name,
                        str(path),
                        file_hash,
                        len(target_paragraphs),
                        json.dumps(metadata, ensure_ascii=False),
                        utc_now(),
                    ),
                )
            for paragraph in source_paragraphs:
                connection.execute(
                    """INSERT OR REPLACE INTO source_paragraphs(
                           source_edition_id, paragraph_index, source_text, source_hash,
                           char_start, char_end
                       ) VALUES(?, ?, ?, ?, ?, ?)""",
                    (
                        edition_id,
                        paragraph.index,
                        paragraph.text,
                        hashlib.sha256(paragraph.text.encode("utf-8")).hexdigest(),
                        paragraph.char_start,
                        paragraph.char_end,
                    ),
                )
            for index, target in enumerate(target_paragraphs[:aligned_count]):
                connection.execute(
                    """INSERT INTO baseline_paragraphs(
                           baseline_document_id, paragraph_index, target_text, target_hash
                       ) VALUES(?, ?, ?, ?)""",
                    (
                        document_id,
                        index,
                        target,
                        hashlib.sha256(target.encode("utf-8")).hexdigest(),
                    ),
                )
            for block, start, alignment_status in block_positions:
                end = start + len(block.source_text)
                ordinal = 0
                for paragraph in source_paragraphs[:aligned_count]:
                    if paragraph.char_end <= start:
                        continue
                    if paragraph.char_start >= end:
                        break
                    overlap_start = max(start, paragraph.char_start)
                    overlap_end = min(end, paragraph.char_end)
                    connection.execute(
                        """INSERT INTO block_baseline_links(
                               block_id, baseline_document_id, paragraph_index, ordinal,
                               overlap_start, overlap_end, partial_start, partial_end,
                               alignment_status
                           ) VALUES(?, ?, ?, ?, ?, ?, ?, ?, ?)""",
                        (
                            block.id,
                            document_id,
                            paragraph.index,
                            ordinal,
                            overlap_start,
                            overlap_end,
                            int(overlap_start > paragraph.char_start),
                            int(overlap_end < paragraph.char_end),
                            alignment_status,
                        ),
                    )
                    ordinal += 1
        return {
            "baseline_document_id": document_id,
            "name": name,
            "source_paragraphs": len(source_paragraphs),
            "target_paragraphs": len(target_paragraphs),
            "aligned_paragraphs": aligned_count,
            "blocks": len(blocks),
            "ambiguous_blocks": ambiguous,
            "idempotent": False,
            "file_sha256": file_hash,
        }

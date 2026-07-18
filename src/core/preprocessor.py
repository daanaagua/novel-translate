"""
预处理模块
负责文本清洗、章节分割、语义分块
"""
import codecs
import hashlib
import re
from dataclasses import dataclass
from typing import List, Optional, Tuple
from pathlib import Path

from ..core.schemas import Book, Chapter, TextChunk
from ..core.epub_reader import EpubReader


@dataclass(frozen=True)
class SourceDocument:
    """已解码的规范源文档及其可审计来源。"""

    text: str
    source_format: str
    encoding: str
    extractor: str
    canonical_segments: tuple[dict, ...]
    excluded_raw_ranges: tuple[dict, ...]
    raw_size: int | None = None
    raw_sha256: str | None = None


_BOM_ENCODINGS = (
    (codecs.BOM_UTF32_LE, "utf-32-le", "UTF32_LE_BOM"),
    (codecs.BOM_UTF32_BE, "utf-32-be", "UTF32_BE_BOM"),
    (codecs.BOM_UTF8, "utf-8-sig", "UTF8_BOM"),
    (codecs.BOM_UTF16_LE, "utf-16-le", "UTF16_LE_BOM"),
    (codecs.BOM_UTF16_BE, "utf-16-be", "UTF16_BE_BOM"),
)


def _normalize_newlines(text: str) -> str:
    return text.replace("\r\n", "\n").replace("\r", "\n")


def _detected_bom(raw: bytes) -> tuple[bytes, str, str] | None:
    for marker, encoding, policy in _BOM_ENCODINGS:
        if raw.startswith(marker):
            return marker, encoding, policy
    return None


def _decode_plain_text(raw: bytes, path: Path) -> tuple[str, str]:
    bom = _detected_bom(raw)
    if bom:
        marker, encoding, _policy = bom
        try:
            return _normalize_newlines(raw[len(marker):].decode(encoding)), encoding
        except UnicodeDecodeError as exc:
            raise ValueError(f"无法按 BOM 指定编码读取文件: {path}") from exc

    for encoding in ('utf-8', 'gbk', 'latin-1'):
        try:
            return _normalize_newlines(raw.decode(encoding)), encoding
        except UnicodeDecodeError:
            continue

    raise ValueError(f"无法读取文件，请检查编码: {path}")


class TextPreprocessor:
    """文本预处理器"""
    
    # 常见章节标题模式
    CHAPTER_PATTERNS = [
        r'^Chapter\s+(\d+)',                    # Chapter 1
        r'^CHAPTER\s+(\d+)',                    # CHAPTER 1
        r'^Chapter\s+([IVXLC]+)',               # Chapter IV (罗马数字)
        r'^CHAPTER\s+([IVXLC]+)',               # CHAPTER IV
        r'^\d+\.',                              # 1.
        r'^第.+章',                              # 中文章节
        r'^([IVXLC]+)\s*$',                     # 纯罗马数字 (如 II, III) - 独占一行
    ]
    
    # 卷/书标题模式 (针对多卷本合集)
    VOLUME_PATTERNS = [
        r'^BOOK\s+(?:\d+|[IVXLC]+|ONE|TWO|THREE|FOUR|FIVE|SIX|SEVEN|EIGHT|NINE|TEN|ELEVEN|TWELVE)\s*$',
        r'^THE\s+SHADOW\s+OF\s+THE\s+TORTURER', # 第一卷
        r'^THE\s+CLAW\s+OF\s+THE\s+CONCILIATOR', # 第二卷
        r'^THE\s+SWORD\s+OF\s+THE\s+LICTOR',     # 第三卷
        r'^THE\s+CITADEL\s+OF\s+THE\s+AUTARCH',  # 第四卷
        r'^THE\s+URTH\s+OF\s+THE\s+NEW\s+SUN',   # 续集
    ]
    
    # 场景分隔符
    SCENE_BREAK_PATTERNS = [
        r'^\s*\*\s*\*\s*\*\s*$',                # * * *
        r'^\s*\*{3,}\s*$',                      # ***
        r'^\s*#\s*#\s*#\s*$',                   # # # #
        r'^\s*-{3,}\s*$',                       # ---
    ]
    
    def __init__(self, max_chunk_tokens: int = 1500, overlap_sentences: int = 2):
        """
        初始化预处理器
        
        Args:
            max_chunk_tokens: 每个 Chunk 的最大 Token 数（粗略估算）
            overlap_sentences: 相邻 Chunk 重叠的句子数
        """
        self.max_chunk_tokens = max_chunk_tokens
        self.overlap_sentences = overlap_sentences
    
    def load_document(self, file_path: str) -> SourceDocument:
        """
        加载源文件，同时保留解码、提取和排除策略。
        
        Args:
            file_path: 文件路径
        
        Returns:
            带 provenance 的规范源文档
        """
        path = Path(file_path)
        
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {file_path}")

        suffix = path.suffix.lower()

        if suffix in {'.txt', '.md'}:
            raw = path.read_bytes()
            text, encoding = _decode_plain_text(raw, path)
            bom = _detected_bom(raw)
            raw_start = len(bom[0]) if bom else 0
            excluded = (
                ({"raw_start": 0, "raw_end": raw_start, "policy": bom[2]},)
                if bom
                else ()
            )
            return SourceDocument(
                text=text,
                source_format=suffix,
                encoding=encoding,
                extractor="plain-text-v1",
                canonical_segments=(
                    {
                        "canonical_start": 0,
                        "canonical_end": len(text),
                        "origin_kind": "decoded_bytes",
                        "origin_ref": f"source/original{suffix}",
                        "raw_start": raw_start,
                        "raw_end": len(raw),
                        "transformation": "decode+newline-normalize",
                    },
                ),
                excluded_raw_ranges=excluded,
                raw_size=len(raw),
                raw_sha256=hashlib.sha256(raw).hexdigest(),
            )

        if suffix == '.epub':
            raw = path.read_bytes()
            text, segments = self._load_epub_with_segments(path)
            if path.read_bytes() != raw:
                raise ValueError("source changed while loading document")
            return SourceDocument(
                text=text,
                source_format=suffix,
                encoding="container",
                extractor="epub-spine-v1",
                canonical_segments=tuple(segments),
                excluded_raw_ranges=(
                    {
                        "raw_start": 0,
                        "raw_end": len(raw),
                        "policy": "EPUB_NON_SPINE_DATA",
                    },
                ),
                raw_size=len(raw),
                raw_sha256=hashlib.sha256(raw).hexdigest(),
            )

        if suffix == '.docx':
            raw = path.read_bytes()
            text, segments = self._load_docx_with_segments(path)
            if path.read_bytes() != raw:
                raise ValueError("source changed while loading document")
            return SourceDocument(
                text=text,
                source_format=suffix,
                encoding="container",
                extractor="docx-paragraph-v1",
                canonical_segments=tuple(segments),
                excluded_raw_ranges=(
                    {
                        "raw_start": 0,
                        "raw_end": len(raw),
                        "policy": "DOCX_NON_DOCUMENT_DATA",
                    },
                ),
                raw_size=len(raw),
                raw_sha256=hashlib.sha256(raw).hexdigest(),
            )

        raise ValueError(
            f"不支持的文件格式: {suffix}；支持 .txt、.md、.docx、.epub"
        )

    def load_text(self, file_path: str) -> str:
        """兼容旧调用方；正式导入应使用 :meth:`load_document`。"""
        return self.load_document(file_path).text

    def _load_docx(self, path: Path) -> str:
        """兼容旧调用方的 DOCX 文本加载。"""
        text, _segments = self._load_docx_with_segments(path)
        return text

    def _load_docx_with_segments(self, path: Path) -> tuple[str, list[dict]]:
        """按文档段落序号提取 DOCX，不丢弃空段落或边界空格。"""
        try:
            import docx
            doc = docx.Document(path)
            pieces: list[str] = []
            segments: list[dict] = []
            cursor = 0
            for index, paragraph in enumerate(doc.paragraphs):
                paragraph_text = _normalize_newlines(paragraph.text)
                piece = ("" if index == 0 else "\n\n") + paragraph_text
                pieces.append(piece)
                end = cursor + len(piece)
                segments.append(
                    {
                        "canonical_start": cursor,
                        "canonical_end": end,
                        "origin_kind": "docx_paragraph",
                        "origin_ref": f"word/document.xml#paragraph={index}",
                        "transformation": "paragraph-text-extract+newline-normalize",
                    }
                )
                cursor = end
            return "".join(pieces), segments
        except ImportError:
            raise ImportError("请安装 python-docx 以支持 .docx 文件: pip install python-docx")
        except Exception as e:
            raise ValueError(f"读取 docx 失败: {e}")

    def _load_epub_with_segments(self, path: Path) -> tuple[str, list[dict]]:
        """按 EPUB spine member 提取规范文本和连续 provenance 区间。"""
        document = EpubReader.read(path)
        pieces: list[str] = []
        segments: list[dict] = []
        cursor = 0
        fallback_index = 1

        for index, section in enumerate(document.sections):
            heading = section.title.strip()
            if heading.isdigit():
                chapter_number = int(heading)
                fallback_index = max(fallback_index, chapter_number + 1)
            else:
                chapter_number = fallback_index
                fallback_index += 1

            body = section.text
            body_lines = body.splitlines()
            if body_lines and body_lines[0].strip().casefold() == heading.casefold():
                body = "\n".join(body_lines[1:]).strip()
            display_heading = "" if heading.isdigit() else heading
            rendered = (
                f"Chapter {chapter_number}\n\n{display_heading}\n\n{body}".strip()
            )
            piece = ("" if index == 0 else "\n\n") + _normalize_newlines(rendered)
            pieces.append(piece)
            end = cursor + len(piece)
            segments.append(
                {
                    "canonical_start": cursor,
                    "canonical_end": end,
                    "origin_kind": "epub_spine_member",
                    "origin_ref": section.source_path,
                    "transformation": "xhtml-text-extract+chapter-marker+newline-normalize",
                }
            )
            cursor = end

        return "".join(pieces), segments

    def _load_plain_text(self, path: Path) -> str:
        """兼容旧调用方的纯文本加载。"""
        content, _encoding = self._load_plain_text_with_encoding(path)
        return content

    def _load_plain_text_with_encoding(self, path: Path) -> tuple[str, str]:
        """以确定性顺序解码，且不用替换字符掩盖解码错误。"""
        return _decode_plain_text(path.read_bytes(), path)
    
    def _clean_text(self, text: str) -> str:
        """仅做无损的换行规范化。"""
        return _normalize_newlines(text)
    
    def split_chapters(self, text: str) -> List[Tuple[str, str, str]]:
        """
        将文本分割为章节 (支持多卷结构 + 智能兜底)
        
        Returns:
            [(chapter_id, chapter_title, chapter_content), ...]
        """
        # 1. 尝试正则切分 Volume
        volume_pattern = '|'.join(f'({p})' for p in self.VOLUME_PATTERNS)
        volume_matches = list(re.finditer(volume_pattern, text, re.MULTILINE | re.IGNORECASE))
        
        volumes = []
        if not volume_matches:
            volumes.append(("v01", text))
        else:
            if volume_matches[0].start() > 0:
                volumes.append(("v00", text[:volume_matches[0].start()]))
            
            for i, match in enumerate(volume_matches):
                vol_id = f"v{i+1:02d}"
                start = match.end()
                end = volume_matches[i+1].start() if i + 1 < len(volume_matches) else len(text)
                volumes.append((vol_id, text[start:end]))

        # 2. 在每个 Volume 内切分 Chapter
        chapters = []
        combined_chapter_pattern = '|'.join(f'({p})' for p in self.CHAPTER_PATTERNS)
        
        for vol_id, vol_content in volumes:
            # 尝试正则切分
            vol_chapters = self._regex_split_chapters(vol_content, combined_chapter_pattern, vol_id)
            
            # 检查切分结果是否合理
            # 如果内容很长（>1万字）但只切出不到 2 章，说明正则可能失效
            if len(vol_content) > 10000 and len(vol_chapters) < 2:
                print(f"[Preprocessor] 卷 {vol_id} 正则切分效果不佳 (仅 {len(vol_chapters)} 章)，启用物理兜底分章...")
                vol_chapters = self._fallback_split_chapters(vol_content, vol_id)
            
            chapters.extend(vol_chapters)
            
        return chapters

    def _regex_split_chapters(self, text: str, pattern: str, vol_id: str) -> List[Tuple[str, str, str]]:
        """正则切分逻辑"""
        matches = list(re.finditer(pattern, text, re.MULTILINE))
        if not matches:
            return [(f"{vol_id}_ch00", f"Volume {vol_id}", text)]
            
        chapters = []
        if matches[0].start() > 0:
            pre_content = text[:matches[0].start()].strip()
            if pre_content:
                chapters.append((f"{vol_id}_pre", "Preamble", pre_content))
                
        id_occurrences = {}
        for i, match in enumerate(matches):
            title = match.group().strip()
            start = match.end()
            end = matches[i+1].start() if i + 1 < len(matches) else len(text)
            content = text[start:end].strip()
            
            if content:
                # 尝试提取数字 ID
                nums = re.findall(r'\d+', title)
                if nums:
                    idx = int(nums[0])
                else:
                    # 罗马数字转阿拉伯数字 (简化版，仅处理 I-X)
                    roman_map = {'I':1, 'II':2, 'III':3, 'IV':4, 'V':5, 'VI':6, 'VII':7, 'VIII':8, 'IX':9, 'X':10}
                    clean_title = title.replace('CHAPTER','').replace('Chapter','').strip().upper()
                    idx = roman_map.get(clean_title, i + 1)
                
                base_id = f"{vol_id}_ch{idx:02d}"
                occurrence = id_occurrences.get(base_id, 0) + 1
                id_occurrences[base_id] = occurrence
                ch_id = base_id if occurrence == 1 else f"{base_id}_{occurrence:02d}"
                chapters.append((ch_id, title, content))
        return chapters

    def _fallback_split_chapters(self, text: str, vol_id: str, chunk_size: int = 5000) -> List[Tuple[str, str, str]]:
        """
        兜底分章逻辑：按字数强制切分
        适用于无任何章节标题的文本
        """
        # 简单按双换行符分割段落
        paragraphs = text.split('\n\n')
        current_chunk = []
        current_len = 0
        chapters = []
        idx = 1
        
        for para in paragraphs:
            current_chunk.append(para)
            current_len += len(para)
            
            if current_len >= chunk_size:
                content = '\n\n'.join(current_chunk)
                chapters.append((f"{vol_id}_auto_{idx:03d}", f"Part {idx}", content))
                current_chunk = []
                current_len = 0
                idx += 1
        
        if current_chunk:
            content = '\n\n'.join(current_chunk)
            chapters.append((f"{vol_id}_auto_{idx:03d}", f"Part {idx}", content))
            
        return chapters

    def split_into_chunks(
        self,
        text: str,
        chapter_id: str
    ) -> List[TextChunk]:
        """
        将章节文本分割为语义完整的 Chunk
        
        核心原则：
        1. 不在句子中间切断
        2. 尽量在段落边界切分
        3. 识别场景分隔符作为优先切分点
        
        Args:
            text: 章节文本
            chapter_id: 章节 ID
        
        Returns:
            TextChunk 列表
        """
        # 按段落分割
        paragraphs = self._split_paragraphs(text)
        
        chunks = []
        current_paragraphs = []
        current_tokens = 0
        chunk_index = 0
        
        for para in paragraphs:
            para_tokens = self._estimate_tokens(para)
            
            # 检查是否是场景分隔符
            is_scene_break = self._is_scene_break(para)
            
            # 判断是否需要切分
            should_split = False
            if is_scene_break and current_paragraphs:
                # 场景分隔符是最佳切分点
                should_split = True
            elif current_tokens + para_tokens > self.max_chunk_tokens and current_paragraphs:
                # 超过 Token 限制
                should_split = True
            
            if should_split:
                # 保存当前 Chunk
                chunk = TextChunk(
                    id=f"{chapter_id}_{chunk_index:03d}",
                    chapter_id=chapter_id,
                    index=chunk_index,
                    source_text='\n\n'.join(current_paragraphs),
                    token_count=current_tokens
                )
                chunks.append(chunk)
                chunk_index += 1
                
                # 重置，保留重叠部分
                if self.overlap_sentences > 0 and current_paragraphs:
                    # 从最后一段取几个句子作为重叠
                    overlap_text = self._get_last_sentences(
                        current_paragraphs[-1],
                        self.overlap_sentences
                    )
                    current_paragraphs = [overlap_text] if overlap_text else []
                    current_tokens = self._estimate_tokens(overlap_text) if overlap_text else 0
                else:
                    current_paragraphs = []
                    current_tokens = 0
            
            # 添加当前段落（除非是场景分隔符）
            if not is_scene_break:
                current_paragraphs.append(para)
                current_tokens += para_tokens
        
        # 保存最后一个 Chunk
        if current_paragraphs:
            chunk = TextChunk(
                id=f"{chapter_id}_{chunk_index:03d}",
                chapter_id=chapter_id,
                index=chunk_index,
                source_text='\n\n'.join(current_paragraphs),
                token_count=current_tokens
            )
            chunks.append(chunk)
        
        return chunks
    
    def _split_paragraphs(self, text: str) -> List[str]:
        """按段落分割文本，对超长段落进行硬切分"""
        # 双换行分割段落
        raw_paragraphs = re.split(r'\n\s*\n', text)
        final_paragraphs = []
        
        # 设定硬切分阈值 (字符数)
        # 英文平均单词长度5，1500 tokens 约等于 6000-8000 字符
        MAX_CHAR_PER_PARA = 3000 
        
        for p in raw_paragraphs:
            p = p.strip()
            if not p:
                continue
                
            if len(p) <= MAX_CHAR_PER_PARA:
                final_paragraphs.append(p)
            else:
                # 暴力按句号切分长段落
                sentences = re.split(r'(?<=[.!?。！？])\s+', p)
                current_chunk = []
                current_len = 0
                
                for sent in sentences:
                    current_chunk.append(sent)
                    current_len += len(sent)
                    if current_len >= MAX_CHAR_PER_PARA:
                        final_paragraphs.append(' '.join(current_chunk))
                        current_chunk = []
                        current_len = 0
                
                if current_chunk:
                    final_paragraphs.append(' '.join(current_chunk))
                    
        return final_paragraphs
    
    def _is_scene_break(self, text: str) -> bool:
        """判断是否是场景分隔符"""
        for pattern in self.SCENE_BREAK_PATTERNS:
            if re.match(pattern, text.strip()):
                return True
        return False
    
    def _estimate_tokens(self, text: str) -> int:
        """
        粗略估算 Token 数
        英文约 1 token/word，中文约 1.5 token/字
        """
        # 简单估算：英文按空格分词，中文按字数
        word_count = len(text.split())
        char_count = len(re.findall(r'[\u4e00-\u9fff]', text))
        return word_count + int(char_count * 1.5)
    
    def _get_last_sentences(self, text: str, n: int) -> str:
        """获取文本最后 n 个句子"""
        # 按句号分割
        sentences = re.split(r'(?<=[.!?。！？])\s+', text)
        if len(sentences) <= n:
            return text
        return ' '.join(sentences[-n:])
    
    def create_book(
        self,
        file_path: str,
        book_id: str,
        title: str
    ) -> Book:
        """
        从文件创建 Book 对象
        
        Args:
            file_path: 源文件路径
            book_id: 书籍 ID
            title: 书名
        
        Returns:
            完整的 Book 对象（包含所有章节和分块）
        """
        # 加载文本
        text = self.load_text(file_path)
        
        # 分割章节
        raw_chapters = self.split_chapters(text)
        
        # 构建 Book
        chapters = []
        # 注意：raw_chapters 现在是 (id, title, content)
        for i, (ch_id, chapter_title, chapter_content) in enumerate(raw_chapters):
            # 分块
            chunks = self.split_into_chunks(chapter_content, ch_id)
            
            chapter = Chapter(
                id=ch_id,
                title=chapter_title,
                index=i,
                source_text=chapter_content,
                chunks=chunks
            )
            chapters.append(chapter)
        
        book = Book(
            id=book_id,
            title=title,
            source_file=file_path,
            chapters=chapters
        )
        
        return book

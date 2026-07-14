import json
import tempfile
import threading
import unittest
from contextlib import closing
from pathlib import Path
from types import SimpleNamespace
from urllib.error import HTTPError
from urllib.request import Request, urlopen

from docx import Document

from src.core.history import TranslationMemory
from src.core.schemas import Chapter, TextChunk
from src.core.v4.baseline import DocxBaselineImporter, paragraph_spans
from src.core.v4.context import ContextBuilder
from src.core.v4.database import V4Database
from src.core.v4.exporter import ParallelV4BookExporter
from src.core.v4.migration import V4Migrator
from src.core.v4.models import ScanOutcome, ScanResponse
from src.core.v4.repairer import V4Repairer
from src.core.v4.verifier import V4Verifier
from src.core.v4.web_review import create_review_server


class FakeVerifyLLM:
    def __init__(self, quote="Archon"):
        self.quote = quote

    def get_model(self, purpose):
        return "fake-verify"

    def chat(self, **kwargs):
        return json.dumps(
            {
                "verdict": "support",
                "rationale": "证据直接支持该译法。",
                "evidence_quotes": [self.quote],
            },
            ensure_ascii=False,
        )


class FakeRepairLLM:
    def __init__(self, paragraphs):
        self.paragraphs = paragraphs

    def get_model(self, purpose):
        return "fake-repair"

    def chat(self, **kwargs):
        return json.dumps(
            {"paragraphs": self.paragraphs, "repair_notes": ["已按问题局部修复"]},
            ensure_ascii=False,
        )


class ParallelV4V2Tests(unittest.TestCase):
    def setUp(self):
        self.temp = tempfile.TemporaryDirectory()
        self.root = Path(self.temp.name) / "book"
        self.root.mkdir()
        self.source = "Alpha begins.\n\nBeta continues for longer.\n\nGamma ends."
        (self.root / "source.txt").write_text(self.source, encoding="utf-8")
        (self.root / "glossary").mkdir()
        memory = TranslationMemory(self.root)
        memory.initialize_chapter(
            Chapter(
                id="ch01",
                title="Chapter One",
                index=0,
                source_text=self.source,
                chunks=[
                    TextChunk(
                        id="ch01_000",
                        chapter_id="ch01",
                        index=0,
                        source_text="Alpha begins.\n\nBeta continues",
                    ),
                    TextChunk(
                        id="ch01_001",
                        chapter_id="ch01",
                        index=1,
                        source_text="Beta continues for longer.\n\nGamma ends.",
                    ),
                ],
            )
        )
        self.project = SimpleNamespace(
            root_dir=self.root,
            source_file=self.root / "source.txt",
            glossary_dir=self.root / "glossary",
            memory=memory,
            book_id="v2-test",
            config_file=self.root / "config.yaml",
        )
        self.project.config_file.write_text("title: V2测试\n", encoding="utf-8")
        V4Migrator(self.project).migrate()
        self.database = V4Database(self.root)

    def tearDown(self):
        self.temp.cleanup()

    @staticmethod
    def write_docx(path: Path, paragraphs):
        document = Document()
        for text in paragraphs:
            document.add_paragraph(text)
        document.save(path)

    def test_paragraph_spans_and_docx_baseline_are_exact_and_idempotent(self):
        spans = paragraph_spans(self.source)
        self.assertEqual([span.text for span in spans], [
            "Alpha begins.",
            "Beta continues for longer.",
            "Gamma ends.",
        ])
        docx = self.root / "baseline.docx"
        self.write_docx(docx, ["阿尔法开始。", "贝塔继续。", "伽马结束。"])
        importer = DocxBaselineImporter(self.database, self.project)
        first = importer.import_docx(docx)
        second = importer.import_docx(docx)
        self.assertFalse(first["idempotent"])
        self.assertTrue(second["idempotent"])
        self.assertEqual(first["aligned_paragraphs"], 3)
        blocks = self.database.list_blocks()
        baseline = self.database.baseline_for_block(blocks[0].id)
        self.assertEqual(baseline["text"], "阿尔法开始。\n\n贝塔继续。")
        self.assertTrue(baseline["has_partial_boundary"])
        self.assertEqual(len(self.database.list_baseline_documents()), 1)

    def test_docx_count_mismatch_is_rejected_without_partial_import(self):
        docx = self.root / "short.docx"
        self.write_docx(docx, ["只有一段。"])
        with self.assertRaisesRegex(ValueError, "段落数不一致"):
            DocxBaselineImporter(self.database, self.project).import_docx(docx)
        self.assertEqual(self.database.list_baseline_documents(), [])

    def test_reveal_boundary_prevents_future_constraint_leakage(self):
        self.database.create_claim(
            kind="translation_constraint",
            statement="保持称呼中性。",
            reveal_global_index=1,
            status="verified",
            locked=True,
        )
        blocks = self.database.list_blocks()
        early = ContextBuilder(self.database, 10000).build(blocks[0])
        later = ContextBuilder(self.database, 10000).build(blocks[1])
        self.assertNotIn("保持称呼中性", early.rendered)
        self.assertIn("保持称呼中性", later.rendered)

    def test_identity_and_time_claims_are_automatically_high_impact(self):
        claim_id = self.database.create_claim(
            kind="identity_hypothesis",
            statement="甲可能是乙。",
            reveal_global_index=1,
        )
        claim = next(item for item in self.database.list_claims() if item["id"] == claim_id)
        self.assertEqual(claim["high_impact"], 1)
        task = self.database.list_verification_tasks()[0]
        self.assertEqual(task["subject_type"], "claim")
        self.assertEqual(task["subject_id"], claim_id)

    def test_scanned_constraint_is_not_injected_before_two_votes(self):
        block = self.database.list_blocks()[0]
        response = ScanResponse.model_validate(
            {
                "mentions": [],
                "ambiguities": [
                    {
                        "paragraph_id": "P000",
                        "evidence_quote": "Alpha begins",
                        "constraint": "保持指代中性。",
                        "confidence": 0.9,
                    }
                ],
            }
        )
        self.database.start_run("scan-constraint", "scan", {})
        self.database.commit_scan_batch(
            "scan-constraint", [ScanOutcome(block=block, response=response)], "fake-scan"
        )
        self.database.finish_run("scan-constraint", "completed")
        self.assertNotIn("保持指代中性", ContextBuilder(self.database, 10000).build(block).rendered)
        result = V4Verifier(
            self.database, lambda: FakeVerifyLLM("Alpha begins")
        ).run()
        self.assertEqual(result["verified"], 1)
        self.assertIn("保持指代中性", ContextBuilder(self.database, 10000).build(block).rendered)

    def test_high_impact_task_requires_two_independent_votes(self):
        concept_id = self.database.import_legacy_concept(
            "Archon", "执政官", "title", "政治职衔"
        )
        with self.database.transaction() as connection:
            connection.execute(
                """INSERT INTO verification_tasks(
                       id, subject_type, subject_id, payload_json, status,
                       required_votes, created_at
                   ) VALUES('verify_test', 'concept', ?, ?, 'open', 2, 'now')""",
                (
                    concept_id,
                    json.dumps(
                        {
                            "source": "Archon",
                            "target": "执政官",
                            "evidence": [{"evidence_quote": "Archon"}],
                        }
                    ),
                ),
            )
        created = []

        def factory():
            created.append(FakeVerifyLLM())
            return created[-1]

        pending = self.database.concepts_for_text("Archon")[0]
        self.assertTrue(pending["verification_pending"])
        self.assertEqual(pending["default_target"], "")
        result = V4Verifier(self.database, factory).run()
        self.assertEqual(result["verified"], 1)
        self.assertEqual(len(created), 2)
        with closing(self.database.connect()) as connection:
            self.assertEqual(
                connection.execute(
                    "SELECT COUNT(*) FROM verification_votes WHERE task_id='verify_test'"
                ).fetchone()[0],
                2,
            )
            status = connection.execute(
                "SELECT status FROM concepts WHERE id=?", (concept_id,)
            ).fetchone()[0]
        self.assertEqual(status, "verified")
        verified = self.database.concepts_for_text("Archon")[0]
        self.assertFalse(verified["verification_pending"])
        self.assertEqual(verified["default_target"], "执政官")

    def test_human_term_lock_invalidates_dependent_parallel_translation(self):
        block = self.database.list_blocks()[0]
        concept_id = self.database.import_legacy_concept(
            "Archon", "执政官", "title", "政治职衔"
        )
        with self.database.transaction() as connection:
            cursor = connection.execute(
                """INSERT INTO translation_versions(
                       block_id, pipeline, knowledge_version, status,
                       draft_translation, final_translation, active, created_at
                   ) VALUES(?, 'parallel_v4', 1, 'completed', '旧稿', '旧稿', 1, 'now')""",
                (block.id,),
            )
            connection.execute(
                """INSERT INTO dependencies(
                       translation_id, dependency_type, dependency_id, knowledge_version
                   ) VALUES(?, 'concept', ?, 1)""",
                (cursor.lastrowid, concept_id),
            )
            connection.execute(
                "UPDATE blocks SET status='completed' WHERE id=?", (block.id,)
            )
        result = self.database.lock_concept_translation(
            "Archon", "阁下", kind="title"
        )
        self.assertEqual(result["affected_translations"], 1)
        concept = self.database.concepts_for_text("Archon")[0]
        self.assertEqual(concept["default_target"], "阁下")
        self.assertTrue(concept["locked"])
        active = self.database.active_translations()[block.id]
        self.assertEqual(active["status"], "needs_revalidate")
        self.assertEqual(
            self.database.get_block_by_identifier(block.id).status,
            "needs_revalidate",
        )

    def insert_translations(self):
        with self.database.transaction() as connection:
            for block in self.database.list_blocks():
                paragraphs = [part.strip() for part in block.source_text.split("\n\n") if part.strip()]
                translation = "\n\n".join(f"这是足够完整的译文段落{index}。" for index, _ in enumerate(paragraphs))
                connection.execute(
                    """INSERT INTO translation_versions(
                           block_id, pipeline, knowledge_version, status,
                           draft_translation, final_translation, active, created_at
                       ) VALUES(?, 'parallel_v4', 1, 'completed', ?, ?, 1, 'now')""",
                    (block.id, translation, translation),
                )
                connection.execute("UPDATE blocks SET status='completed' WHERE id=?", (block.id,))

    def test_local_repair_preserves_previous_translation_version(self):
        self.insert_translations()
        block = self.database.list_blocks()[0]
        before = self.database.active_translations()[block.id]["final_translation"]
        result = V4Repairer(
            self.database,
            lambda: FakeRepairLLM(
                ["阿尔法由此开始，细节保持完整。", "贝塔继续向前发展，信息没有删节。"]
            ),
        ).run(block_identifier=block.id, issues=["修复专名"])
        self.assertEqual(result["completed"], 1)
        with closing(self.database.connect()) as connection:
            versions = connection.execute(
                """SELECT final_translation, active FROM translation_versions
                   WHERE block_id=? AND pipeline='parallel_v4' ORDER BY id""",
                (block.id,),
            ).fetchall()
            task = connection.execute(
                "SELECT status, translation_id FROM repair_tasks WHERE block_id=?",
                (block.id,),
            ).fetchone()
        self.assertEqual(len(versions), 2)
        self.assertEqual(versions[0]["final_translation"], before)
        self.assertEqual(versions[0]["active"], 0)
        self.assertEqual(versions[1]["active"], 1)
        self.assertEqual(task["status"], "completed")
        self.assertIsNotNone(task["translation_id"])

    def test_failed_repair_keeps_current_translation_active(self):
        self.insert_translations()
        block = self.database.list_blocks()[0]
        before = self.database.active_translations()[block.id]["final_translation"]
        result = V4Repairer(
            self.database,
            lambda: FakeRepairLLM(["段落数错误。"]),
            max_attempts=1,
        ).run(block_identifier=block.id, issues=["测试失败路径"])
        self.assertEqual(result["needs_human"], 1)
        self.assertEqual(
            self.database.active_translations()[block.id]["final_translation"], before
        )
        self.assertEqual(self.database.list_human_queue()[0]["kind"], "repair_failed")

    def test_export_annotations_are_opt_in_and_approved_only(self):
        self.insert_translations()
        block = self.database.list_blocks()[0]
        approved = self.database.add_annotation(block.id, 0, "批准的注释")
        self.database.resolve_annotation(approved, "approve")
        self.database.add_annotation(block.id, 0, "未批准的注释")
        exporter = ParallelV4BookExporter(self.project, self.database)
        plain = exporter.export_v4(self.root / "plain")
        annotated = exporter.export_v4(
            self.root / "annotated", include_annotations=True
        )
        plain_text = plain.txt_path.read_text(encoding="utf-8-sig")
        annotated_text = annotated.txt_path.read_text(encoding="utf-8-sig")
        self.assertNotIn("批准的注释", plain_text)
        self.assertIn("〔注1〕", annotated_text)
        self.assertIn("批准的注释", annotated_text)
        self.assertNotIn("未批准的注释", annotated_text)

    def test_web_server_is_loopback_and_post_requires_token(self):
        docx = self.root / "web-baseline.docx"
        self.write_docx(docx, ["阿尔法开始。", "贝塔继续。", "伽马结束。"])
        DocxBaselineImporter(self.database, self.project).import_docx(
            docx, name="web_baseline"
        )
        server = create_review_server(self.database, port=0)
        thread = threading.Thread(target=server.serve_forever, daemon=True)
        thread.start()
        try:
            host, port = server.server_address
            self.assertEqual(host, "127.0.0.1")
            state = json.loads(urlopen(f"http://{host}:{port}/api/state").read())
            self.assertIn("blocks", state)
            block = self.database.list_blocks()[0]
            detail = json.loads(
                urlopen(
                    f"http://{host}:{port}/api/block?id={block.id}&blind=1"
                ).read()
            )
            self.assertFalse(detail["blind_available"])
            self.assertEqual(detail["candidates"][0]["origin"], "web_baseline")
            self.assertEqual(
                detail["candidates"][1]["origin"], "parallel_v4（尚未生成）"
            )
            payload = json.dumps(
                {"block": block.id, "paragraph_index": 0, "body": "测试注释"}
            ).encode()
            with self.assertRaises(HTTPError) as denied:
                urlopen(
                    Request(
                        f"http://{host}:{port}/api/annotations",
                        data=payload,
                        headers={"Content-Type": "application/json"},
                        method="POST",
                    )
                )
            self.assertEqual(denied.exception.code, 403)
            denied.exception.close()
            accepted = urlopen(
                Request(
                    f"http://{host}:{port}/api/annotations",
                    data=payload,
                    headers={
                        "Content-Type": "application/json",
                        "X-Review-Token": server.review_token,
                    },
                    method="POST",
                )
            )
            self.assertEqual(accepted.status, 201)
        finally:
            server.shutdown()
            server.server_close()
            thread.join(timeout=2)

    def test_blind_comparison_vote_records_hidden_candidate_origin(self):
        docx = self.root / "vote-baseline.docx"
        self.write_docx(docx, ["阿尔法开始。", "贝塔继续。", "伽马结束。"])
        DocxBaselineImporter(self.database, self.project).import_docx(
            docx, name="vote_baseline"
        )
        with self.database.transaction() as connection:
            connection.execute(
                """UPDATE baseline_paragraphs
                   SET target_text='「''Alpha done.''」'
                   WHERE paragraph_index=0"""
            )
        self.insert_translations()
        server = create_review_server(self.database, port=0)
        thread = threading.Thread(target=server.serve_forever, daemon=True)
        thread.start()
        try:
            host, port = server.server_address
            page = urlopen(f"http://{host}:{port}/").read().decode("utf-8")
            self.assertIn('oninput="saveVoteDraft()"', page)
            self.assertIn("localStorage.setItem(DRAFT_KEY", page)
            self.assertIn("保存理由/更新评价", page)
            self.assertIn("saveCurrentVote()", page)
            block = self.database.list_blocks()[0]
            blind = json.loads(
                urlopen(
                    f"http://{host}:{port}/api/block?id={block.id}&blind=1"
                ).read()
            )
            revealed = json.loads(
                urlopen(
                    f"http://{host}:{port}/api/block?id={block.id}&blind=0"
                ).read()
            )
            self.assertTrue(blind["blind_available"])
            self.assertTrue(all("origin" not in item for item in blind["candidates"]))
            baseline_candidate = next(
                item for item in revealed["candidates"]
                if item["origin"] == "vote_baseline"
            )
            self.assertIn("“‘Alpha done.’”", baseline_candidate["text"])
            self.assertNotIn("「", baseline_candidate["text"])
            payload = json.dumps(
                {
                    "block": block.id,
                    "choice": "A",
                    "blinded": True,
                    "note": "A更流畅",
                }
            ).encode()
            response = urlopen(
                Request(
                    f"http://{host}:{port}/api/comparison-votes",
                    data=payload,
                    headers={
                        "Content-Type": "application/json",
                        "X-Review-Token": server.review_token,
                    },
                    method="POST",
                )
            )
            submitted = json.loads(response.read())
            self.assertEqual(submitted["choice"], "A")
            self.assertNotIn("selected_origin", submitted)
            stored = self.database.comparison_vote_for_block(block.id)
            self.assertEqual(stored["choice"], "A")
            self.assertEqual(stored["candidate_a_origin"], revealed["candidates"][0]["origin"])
            self.assertEqual(stored["selected_origin"], revealed["candidates"][0]["origin"])
            after = json.loads(
                urlopen(
                    f"http://{host}:{port}/api/block?id={block.id}&blind=1"
                ).read()
            )
            self.assertEqual(after["comparison_vote"]["choice"], "A")
            self.assertNotIn("selected_origin", after["comparison_vote"])
            with self.database.transaction() as connection:
                connection.execute(
                    """UPDATE translation_versions
                       SET final_translation=final_translation || ' 修订。'
                       WHERE block_id=? AND pipeline='parallel_v4' AND active=1""",
                    (block.id,),
                )
            stale = json.loads(
                urlopen(
                    f"http://{host}:{port}/api/block?id={block.id}&blind=1"
                ).read()
            )
            self.assertIsNone(stale["comparison_vote"])
            revised_payload = json.dumps(
                {
                    "block": block.id,
                    "choice": "B",
                    "blinded": True,
                    "note": "修订后重新评价",
                }
            ).encode()
            urlopen(
                Request(
                    f"http://{host}:{port}/api/comparison-votes",
                    data=revised_payload,
                    headers={
                        "Content-Type": "application/json",
                        "X-Review-Token": server.review_token,
                    },
                    method="POST",
                )
            ).read()
            self.assertEqual(len(self.database.list_comparison_vote_history()), 1)
            revised_vote = self.database.comparison_vote_for_block(block.id)
            self.assertEqual(revised_vote["choice"], "B")
            self.assertTrue(revised_vote["candidate_a_hash"])
            self.assertTrue(revised_vote["candidate_b_hash"])
        finally:
            server.shutdown()
            server.server_close()
            thread.join(timeout=2)


if __name__ == "__main__":
    unittest.main()
